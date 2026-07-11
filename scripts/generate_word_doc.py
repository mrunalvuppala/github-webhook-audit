"""Generate the AgentAuditAI Word documentation with visual diagrams."""

from __future__ import annotations

import base64
import io
import zlib
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
DESKTOP = Path.home() / "Desktop"
OUTPUT = DOCS_DIR / "AgentAuditAI_Project_Documentation.docx"
DESKTOP_OUTPUT = DESKTOP / "AgentAuditAI_Project_Documentation.docx"
DIAGRAMS_DIR = DOCS_DIR / "diagrams"


def _save_fig(name: str) -> Path:
    path = DIAGRAMS_DIR / name
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def draw_system_architecture() -> Path:
    fig, ax = plt.subplots(figsize=(12, 7.5))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("AgentAudit AI — Multi-Tenant System Architecture", fontsize=16, fontweight="bold", pad=16)

    boxes = [
        (0.4, 6.0, 2.2, 1.0, "#E8F1FF", "GitHub /\nGitHub Enterprise"),
        (3.0, 6.0, 2.5, 1.0, "#DFF5E8", "FastAPI Web\nHMAC + 202 ACK"),
        (5.9, 6.0, 1.8, 1.0, "#FFF2CC", "Redis\nCelery Broker"),
        (8.1, 6.0, 2.0, 1.0, "#FCE4EC", "Celery Worker\nSecurityEngine"),
        (3.0, 3.5, 2.5, 1.0, "#E8EAF6", "Demo UI\nGET / + /v1/demo/audit"),
        (5.9, 3.5, 2.0, 1.0, "#BBDEFB", "PostgreSQL 16\norganizations"),
        (8.4, 3.5, 1.7, 1.0, "#C5CAE9", "scan_audit_logs\nRLS enforced"),
        (10.3, 3.5, 1.4, 1.0, "#E0F7FA", "AST + Secrets\nAWS/Stripe/GitLab"),
    ]

    for x, y, w, h, color, label in boxes:
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.05,rounding_size=0.08",
            linewidth=1.2, edgecolor="#455A64", facecolor=color
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=8.5, fontweight="bold")

    arrows = [
        ((2.6, 6.5), (3.0, 6.5), "Signed webhook"),
        ((5.5, 6.5), (5.9, 6.5), "Queue task"),
        ((7.7, 6.5), (8.1, 6.5), "Process"),
        ((9.1, 6.0), (9.1, 4.5), "SET LOCAL tenant"),
        ((9.1, 4.0), (7.9, 4.0), "Persist metadata"),
        ((4.25, 4.5), (4.25, 6.0), "Dev audit"),
        ((3.0, 6.2), (2.6, 6.2), "403 / 202"),
    ]
    for start, end, label in arrows:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=1.5, color="#37474F"))
        mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        ax.text(mx, my + 0.15, label, fontsize=7.5, ha="center", color="#37474F")

    return _save_fig("01_system_architecture.png")


def draw_request_lifecycle() -> Path:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    ax.axis("off")
    ax.set_title("Production Webhook Request Lifecycle", fontsize=15, fontweight="bold", pad=14)

    steps = [
        "1. GitHub POST\n/v1/webhooks/github",
        "2. Read raw body\nbytes",
        "3. Verify HMAC\nSHA-256",
        "4. Extract\ninstallation_id",
        "5. Queue Celery\ntask (agentaudit)",
        "6. Return 202\nAccepted",
        "7. SET LOCAL\ntenant context",
        "8. Scan diff +\npersist RLS row",
    ]
    x_positions = [0.4 + i * 1.35 for i in range(len(steps))]
    for i, (x, step) in enumerate(zip(x_positions, steps)):
        color = "#C8E6C9" if i not in (2, 7) else "#FFCDD2" if i == 2 else "#BBDEFB"
        rect = mpatches.FancyBboxPatch(
            (x, 2.0), 1.15, 1.35, boxstyle="round,pad=0.04,rounding_size=0.08",
            linewidth=1.1, edgecolor="#455A64", facecolor=color
        )
        ax.add_patch(rect)
        ax.text(x + 0.58, 2.68, step, ha="center", va="center", fontsize=8.5, fontweight="bold")
        if i < len(steps) - 1:
            ax.annotate("", xy=(x + 1.2, 2.68), xytext=(x + 1.15, 2.68),
                        arrowprops=dict(arrowstyle="->", lw=1.4, color="#37474F"))

    ax.text(6.0, 0.8, "Invalid signature path returns 403 immediately — no diff processing occurs",
            ha="center", fontsize=10, color="#B71C1C", style="italic")
    ax.set_xlim(0, 11.5)
    ax.set_ylim(0, 4.5)
    return _save_fig("02_request_lifecycle.png")


def draw_database_rls() -> Path:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    ax.axis("off")
    ax.set_title("PostgreSQL Row-Level Security — Tenant Isolation", fontsize=15, fontweight="bold", pad=14)

    boxes = [
        (0.5, 3.5, 2.5, 1.2, "#E3F2FD", "Celery Worker\nTransaction Start"),
        (3.5, 3.5, 3.2, 1.2, "#FFF9C4", "SET LOCAL\napp.current_organization_id"),
        (7.2, 3.5, 2.5, 1.2, "#C8E6C9", "INSERT\nscan_audit_logs"),
        (10.2, 3.5, 1.5, 1.2, "#FFCDD2", "Other Tenant\n0 rows visible"),
        (3.5, 1.5, 3.2, 1.0, "#D1C4E9", "RLS Policy:\norganization_id = current_setting(...)"),
        (7.2, 1.5, 2.5, 1.0, "#BBDEFB", "organizations\ninstallation mapping"),
    ]
    for x, y, w, h, color, label in boxes:
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.04,rounding_size=0.08",
            linewidth=1.1, edgecolor="#455A64", facecolor=color
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=8.5, fontweight="bold")

    for start, end in [((3.0, 4.1), (3.5, 4.1)), ((6.7, 4.1), (7.2, 4.1)), ((9.7, 4.1), (10.2, 4.1)),
                       ((5.1, 3.5), (5.1, 2.5)), ((8.45, 3.5), (8.45, 2.5))]:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=1.4, color="#37474F"))

    ax.set_xlim(0, 12)
    ax.set_ylim(1.0, 5.5)
    return _save_fig("07_database_rls.png")


def draw_enterprise_deployment() -> Path:
    fig, ax = plt.subplots(figsize=(12, 6))
    ax.axis("off")
    ax.set_title("Enterprise Integration — GitHub Enterprise Server", fontsize=15, fontweight="bold", pad=14)

    layers = [
        (0.5, 5.0, 11.0, 0.9, "#E3F2FD", "Corporate Network / VPC"),
        (1.0, 3.7, 3.0, 1.0, "#FFF9C4", "GitHub\nEnterprise Server"),
        (4.5, 3.7, 2.5, 1.0, "#C8E6C9", "Internal\nLoad Balancer"),
        (7.5, 3.7, 3.5, 1.0, "#FFE0B2", "AgentAudit AI\nWeb + Workers"),
        (1.0, 1.8, 2.3, 1.0, "#F8BBD0", "Enterprise\nRedis"),
        (3.8, 1.8, 2.8, 1.0, "#D1C4E9", "PostgreSQL\nRLS Tenant DB"),
        (7.0, 1.8, 2.5, 1.0, "#B2DFDB", "HashiCorp\nVault"),
        (9.8, 1.8, 2.0, 1.0, "#CFD8DC", "SIEM\nSplunk/Datadog"),
    ]
    for x, y, w, h, color, label in layers:
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.04,rounding_size=0.06",
            linewidth=1.1, edgecolor="#455A64", facecolor=color
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=9, fontweight="bold")

    flows = [((4.0, 4.2), (4.5, 4.2)), ((7.0, 4.2), (7.5, 4.2)), ((8.5, 3.7), (2.1, 2.8)),
             ((8.5, 3.7), (5.2, 2.8)), ((8.5, 3.7), (8.2, 2.8)), ((10.5, 3.7), (10.8, 2.8))]
    for start, end in flows:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=1.3, color="#37474F"))

    ax.set_xlim(0, 12.5)
    ax.set_ylim(1.0, 6.2)
    return _save_fig("03_enterprise_deployment.png")


def draw_public_company_deployment() -> Path:
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.axis("off")
    ax.set_title("Public Company Integration — SOX / SEC / GDPR Alignment", fontsize=15, fontweight="bold", pad=14)

    zones = [
        (0.4, 5.5, 11.5, 0.8, "#ECEFF1", "Internet Zone"),
        (0.4, 4.2, 11.5, 0.8, "#FFECB3", "DMZ — Cloud WAF + TLS Termination"),
        (0.4, 1.0, 11.5, 2.8, "#E8F5E9", "Private Cloud — AgentAuditAI Platform"),
        (0.4, 0.2, 11.5, 0.6, "#E1F5FE", "GRC Zone — SIEM, Legal Hold, Board Reporting"),
    ]
    for x, y, w, h, color, label in zones:
        rect = mpatches.Rectangle((x, y), w, h, linewidth=1.0, edgecolor="#607D8B", facecolor=color, alpha=0.55)
        ax.add_patch(rect)
        ax.text(x + 0.15, y + h - 0.18, label, fontsize=9, fontweight="bold", va="top")

    nodes = [
        (1.0, 5.7, 2.0, 0.5, "GitHub Cloud"),
        (5.0, 4.35, 2.5, 0.5, "Cloud WAF"),
        (1.2, 2.5, 2.0, 0.8, "API Gateway"),
        (3.8, 2.5, 1.8, 0.8, "Redis"),
        (6.0, 2.5, 2.0, 0.8, "Celery Workers"),
        (8.5, 2.5, 2.3, 0.8, "Audit Metadata DB"),
        (1.5, 0.45, 2.0, 0.45, "SIEM"),
        (4.5, 0.45, 2.2, 0.45, "GRC Platform"),
        (8.0, 0.45, 2.5, 0.45, "Legal Hold Archive"),
    ]
    for x, y, w, h, label in nodes:
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.03,rounding_size=0.05",
            linewidth=1.0, edgecolor="#455A64", facecolor="white"
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=8.5, fontweight="bold")

    for start, end in [((3.0, 5.95), (5.0, 4.6)), ((7.5, 4.6), (2.2, 3.3)), ((3.2, 2.9), (3.8, 2.9)),
                       ((5.6, 2.9), (6.0, 2.9)), ((8.0, 2.9), (8.5, 2.9)), ((7.2, 2.5), (2.5, 0.9)),
                       ((9.5, 2.5), (5.6, 0.9)), ((7.0, 2.5), (9.2, 0.9))]:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=1.2, color="#37474F"))

    ax.set_xlim(0, 12.5)
    ax.set_ylim(0, 6.8)
    return _save_fig("04_public_company_deployment.png")


def draw_ui_layout() -> Path:
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.axis("off")
    ax.set_title("Demo UI Dashboard Layout", fontsize=15, fontweight="bold", pad=14)

    ax.add_patch(mpatches.Rectangle((0.5, 5.2), 10.0, 0.8, facecolor="#1A237E", edgecolor="none"))
    ax.text(5.5, 5.6, "AgentAuditAI — Live Demo Dashboard", color="white", ha="center", va="center", fontsize=13, fontweight="bold")
    ax.text(9.8, 5.6, "API Online", color="#A5D6A7", ha="center", va="center", fontsize=9, fontweight="bold")

    ax.add_patch(mpatches.Rectangle((0.5, 0.8), 5.8, 4.0, facecolor="#ECEFF1", edgecolor="#607D8B"))
    ax.text(3.4, 4.45, "Diff Input Panel", ha="center", fontsize=11, fontweight="bold")
    ax.text(1.0, 4.0, "[Clean code] [AWS leak] [Stripe leak] [GitLab token]", fontsize=8.5)
    ax.add_patch(mpatches.Rectangle((1.0, 2.0), 4.8, 1.7, facecolor="#263238", edgecolor="#455A64"))
    ax.text(3.4, 2.85, "+ AWS_KEY = 'AKIA...'\n+ stripe_key = 'sk_live_...'", color="#ECEFF1", ha="center", va="center", fontsize=9, family="monospace")
    ax.add_patch(mpatches.Rectangle((1.0, 1.2), 2.0, 0.5, facecolor="#1565C0", edgecolor="none"))
    ax.text(2.0, 1.45, "Run audit", color="white", ha="center", va="center", fontsize=9, fontweight="bold")
    ax.add_patch(mpatches.Rectangle((3.2, 1.2), 2.6, 0.5, facecolor="#455A64", edgecolor="none"))
    ax.text(4.5, 1.45, "Queue via webhook", color="white", ha="center", va="center", fontsize=8.5, fontweight="bold")

    ax.add_patch(mpatches.Rectangle((6.7, 0.8), 3.8, 4.0, facecolor="#ECEFF1", edgecolor="#607D8B"))
    ax.text(8.6, 4.45, "Audit Result Panel", ha="center", fontsize=11, fontweight="bold")
    ax.add_patch(mpatches.Rectangle((7.2, 2.0), 3.0, 2.0, facecolor="#263238", edgecolor="#455A64"))
    ax.text(8.7, 3.2, "FAIL\nHigh risk: YES\nViolations: 1\naws_access_key_id (high)", color="#EF9A9A", ha="center", va="center", fontsize=9, family="monospace")

    ax.set_xlim(0, 11.5)
    ax.set_ylim(0, 6.5)
    return _save_fig("05_ui_dashboard.png")


def draw_testing_flow() -> Path:
    fig, ax = plt.subplots(figsize=(11, 5))
    ax.axis("off")
    ax.set_title("Testing Workflow — Recommended Sequence", fontsize=15, fontweight="bold", pad=14)

    phases = [
        ("Step 1\nStart Services", "docker-compose up -d", "#BBDEFB"),
        ("Step 2\nHealth Check", "GET /health → ok", "#C8E6C9"),
        ("Step 3\nUI Audit", "Run audit → PASS/FAIL", "#FFE082"),
        ("Step 4\nWebhook Test", "Signed POST → 202", "#CE93D8"),
        ("Step 5\nWorker Verify", "Check worker + DB logs", "#90CAF9"),
    ]
    for i, (title, detail, color) in enumerate(phases):
        x = 0.5 + i * 2.15
        rect = mpatches.FancyBboxPatch((x, 2.0), 1.9, 1.8, boxstyle="round,pad=0.05", facecolor=color, edgecolor="#455A64")
        ax.add_patch(rect)
        ax.text(x + 0.95, 3.2, title, ha="center", va="center", fontsize=9, fontweight="bold")
        ax.text(x + 0.95, 2.5, detail, ha="center", va="center", fontsize=8)
        if i < len(phases) - 1:
            ax.annotate("", xy=(x + 2.0, 2.9), xytext=(x + 1.9, 2.9), arrowprops=dict(arrowstyle="->", lw=1.4))

    ax.set_xlim(0, 11.5)
    ax.set_ylim(1.5, 4.5)
    return _save_fig("06_testing_flow.png")


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_color: str = "1A237E") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cells[i].text = text
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(26, 35, 126)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_image_page(doc: Document, image_path: Path, caption: str) -> None:
    doc.add_page_break()
    add_heading(doc, caption, level=2)
    doc.add_picture(str(image_path), width=Inches(6.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(96, 125, 139)


def build_document(images: dict[str, Path]) -> Document:
    doc = Document()

    # Cover page
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AgentAuditAI")
    run.bold = True
    run.font.size = Pt(34)
    run.font.color.rgb = RGBColor(26, 35, 126)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Complete Project & Technical Documentation")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(55, 71, 79)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("\nVersion 2.0.0\n").bold = True
    meta.add_run("Author: Naga Sai Mrunal Vuppala\n")
    meta.add_run("Architecture designed, engineered, and maintained by Naga Sai Mrunal Vuppala\n")
    meta.add_run("Repository: github.com/mrunalvuppala/github-webhook-audit\n")
    meta.add_run("Date: July 2026\n")

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag.add_run("\nEnterprise • Public Company • Compliance Ready")
    tag_run.italic = True
    tag_run.font.color.rgb = RGBColor(21, 101, 192)

    doc.add_page_break()

    # Table of contents
    add_heading(doc, "Table of Contents", 1)
    toc = [
        "1. Executive Summary",
        "2. System Architecture (Visual)",
        "3. Component Reference",
        "4. Security & Compliance Design",
        "5. Installation & Operations",
        "6. Testing Guide (Visual)",
        "7. Enterprise Integration (Visual)",
        "8. Public Company Integration (Visual)",
        "9. API Reference",
        "10. Troubleshooting",
        "11. Roadmap & Extensions",
        "12. Technology Stack & Rationale",
        "Appendix A — Environment Template",
        "Appendix B — Support",
    ]
    for item in toc:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # Section 1
    add_heading(doc, "1. Executive Summary", 1)
    add_body(doc, "AgentAudit AI is a production-grade, multi-tenant GitHub webhook security platform with PostgreSQL Row-Level Security, Redis-backed Celery workers, and AST/secret scanning.")
    add_heading(doc, "Key Capabilities", 2)
    add_table(doc, ["Capability", "Description"], [
        ["Webhook verification", "HMAC-SHA256 via X-Hub-Signature-256"],
        ["Asynchronous auditing", "Celery workers with tenant-scoped DB writes"],
        ["Credential detection", "AWS, Stripe, GitLab token patterns"],
        ["AST validation", "eval(), exec(), os.system(), subprocess.*"],
        ["Multi-tenant isolation", "PostgreSQL RLS on scan_audit_logs"],
        ["Demo UI", "Browser dashboard for PASS/FAIL presentations"],
    ])
    add_heading(doc, "Technology Stack", 2)
    add_table(doc, ["Layer", "Technology"], [
        ["API Gateway", "FastAPI + Uvicorn"],
        ["Configuration", "Pydantic Settings v2 (app/config.py)"],
        ["Task Queue", "Celery + Redis"],
        ["Database", "PostgreSQL 16 + SQLAlchemy 2.0"],
        ["Scan Engine", "SecurityEngine + ASTParser"],
        ["Deployment", "Docker Compose (postgres, redis, web, worker)"],
        ["License", "Business Source License 1.1"],
    ])

    add_image_page(doc, images["architecture"], "Figure 1 — System Architecture Overview")

    add_heading(doc, "Request Lifecycle", 2)
    add_body(doc, "Production webhooks follow an eight-step lifecycle ending in metadata-only persistence. Invalid signatures return HTTP 403 before any diff processing.")
    add_image_page(doc, images["lifecycle"], "Figure 2 — Production Webhook Request Lifecycle")
    add_image_page(doc, images["rls"], "Figure 3 — PostgreSQL Row-Level Security Tenant Isolation")

    # Section 3
    doc.add_page_break()
    add_heading(doc, "3. Component Reference", 1)
    add_heading(doc, "Configuration Variables", 2)
    add_table(doc, ["Variable", "Required", "Default", "Purpose"], [
        ["DATABASE_URL", "Yes", "—", "PostgreSQL connection string"],
        ["REDIS_URL", "No", "redis://localhost:6379/0", "Celery broker/backend"],
        ["GITHUB_WEBHOOK_SECRET", "Yes", "—", "Webhook HMAC secret"],
        ["SECRET_KEY", "Yes", "—", "Application cryptographic secret"],
        ["SECURITY_ENVIRONMENT", "No", "development", "Enables demo endpoints when development"],
    ])
    add_heading(doc, "Detection Rules", 2)
    add_table(doc, ["Rule ID", "Pattern", "Risk"], [
        ["aws_access_key_id", "AKIA[0-9A-Z]{16}", "High"],
        ["stripe_live_secret_key", "sk_live_* (24+ chars)", "High"],
        ["gitlab_access_token", "glpat-*", "High"],
        ["dangerous_call", "eval(), exec(), os.system()", "Medium"],
        ["subprocess_call", "subprocess.run(), subprocess.call()", "Medium"],
    ])
    add_heading(doc, "API Endpoints", 2)
    add_table(doc, ["Endpoint", "Method", "Purpose"], [
        ["/", "GET", "Demo UI dashboard"],
        ["/health", "GET", "Health check (DB + Redis)"],
        ["/docs", "GET", "Swagger UI"],
        ["/v1/webhooks/github", "POST", "Production webhook ingress"],
        ["/v1/scan", "POST", "Synchronous AST + secret scan"],
        ["/v1/demo/audit", "POST", "Dev-only synchronous audit"],
        ["/v1/demo/webhook", "POST", "Dev-only Celery queue test"],
    ])

    # Section 4
    doc.add_page_break()
    add_heading(doc, "4. Security & Compliance Design", 1)
    add_table(doc, ["Requirement", "AgentAudit AI Response"], [
        ["Data minimization", "Only audit metadata stored; never raw diffs"],
        ["Immediate webhook ACK", "202 response prevents GitHub retry storms"],
        ["Secret verification", "Rejects unsigned/tampered payloads (403)"],
        ["Memory hygiene", "Explicit diff purge + garbage collection"],
        ["Tenant isolation", "PostgreSQL RLS + SET LOCAL tenant context"],
        ["Audit trail", "scan_audit_logs with JSONB findings"],
    ])
    add_heading(doc, "Production Hardening Checklist", 2)
    checklist = [
        "Set SECURITY_ENVIRONMENT=production",
        "Use managed Redis with TLS",
        "Store secrets in Vault (AWS/Azure/HashiCorp)",
        "Disable demo endpoints in production",
        "Enable HTTPS at load balancer",
        "Restrict ingress to GitHub IP ranges",
        "Connect production DATABASE_URL with RLS enabled",
        "Enable centralized logging (Splunk/Datadog/ELK)",
        "Rotate GITHUB_WEBHOOK_SECRET on schedule",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")
    add_heading(doc, "Never Logged", 2)
    for item in ["Raw webhook body", "Diff content", "Matched secret values", "GITHUB_WEBHOOK_SECRET"]:
        doc.add_paragraph(item, style="List Bullet")

    # Section 5
    doc.add_page_break()
    add_heading(doc, "5. Installation & Operations", 1)
    add_body(doc, "Recommended quick start:")
    code = doc.add_paragraph("cd C:\\git\\github-webhook-audit\ncopy .env.example .env\nstart.bat")
    for run in code.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    add_table(doc, ["Service", "Port", "Role"], [
        ["postgres", "5432", "Multi-tenant metadata store"],
        ["redis", "6379", "Celery broker"],
        ["web", "8000", "FastAPI gateway + UI"],
        ["worker", "—", "Background audits"],
    ])

    add_image_page(doc, images["ui"], "Figure 4 — Demo UI Dashboard Layout")

    # Section 6
    doc.add_page_break()
    add_heading(doc, "6. Testing Guide", 1)
    add_table(doc, ["Test", "Method", "Expected"], [
        ["Health check", "GET /health", "200 OK"],
        ["UI dashboard", "GET /", "200 HTML page"],
        ["Clean diff audit", "UI → Run audit", "PASS, 0 violations"],
        ["AWS key leak", "UI → AWS example", "FAIL, high risk"],
        ["Invalid signature", "Webhook without HMAC", "403 Forbidden"],
        ["Async pipeline", "Queue via webhook flow", "202 + worker logs"],
    ])
    add_image_page(doc, images["testing"], "Figure 5 — Recommended Testing Workflow")

    add_heading(doc, "Quick Test Example", 2)
    add_body(doc, "Run: python scripts\\test_scan_engine.py")
    add_body(doc, "Expected output: secret test blocked, AST test blocked, clean test ok.")

    # Section 7
    doc.add_page_break()
    add_heading(doc, "7. Enterprise Integration", 1)
    add_body(doc, "Deploy inside corporate VPC for GitHub Enterprise Server. Route webhooks to https://audit.internal.company.com/v1/webhooks/github")
    add_image_page(doc, images["enterprise"], "Figure 6 — Enterprise Deployment Architecture")
    add_table(doc, ["Control", "Recommendation"], [
        ["Network", "Private subnet behind WAF"],
        ["Secrets", "Vault-backed secret injection"],
        ["TLS", "Terminate at ALB/NGINX with corporate CA"],
        ["IAM", "Service accounts for worker → DB access"],
        ["Audit", "Forward metadata logs to SIEM"],
        ["HA", "Multiple Celery workers + Redis Sentinel"],
    ])

    # Section 8
    doc.add_page_break()
    add_heading(doc, "8. Public Company Integration", 1)
    add_table(doc, ["Regulation", "Alignment"], [
        ["SOX ITGC", "Prevent unauthorized credential commits; audit trail"],
        ["SEC Cybersecurity Disclosure", "Proactive secret scanning controls"],
        ["GDPR Art. 5", "Data minimization — no diff retention"],
        ["PCI-DSS", "Prevent Stripe key leakage in source code"],
        ["NIST CSF", "Detect (DE.CM), Respond (RS.AN)"],
    ])
    add_image_page(doc, images["public"], "Figure 7 — Public Company Deployment Pattern")
    add_table(doc, ["Metric", "Business Value"], [
        ["Total audits per quarter", "Control operating effectiveness"],
        ["High-risk detection rate", "Security posture trend"],
        ["Mean time to detect", "Incident response KPI"],
        ["Violations by tenant/repo", "Targeted developer training"],
        ["False positive rate", "Rule tuning effectiveness"],
    ])

    # Section 9
    doc.add_page_break()
    add_heading(doc, "9. API Reference", 1)
    add_heading(doc, "POST /v1/webhooks/github", 2)
    add_body(doc, "Headers: X-Hub-Signature-256, Content-Type: application/json")
    add_body(doc, "Responses: 202 Accepted | 403 Forbidden | 400 Bad Request")

    # Section 10
    add_heading(doc, "10. Troubleshooting", 1)
    add_table(doc, ["Symptom", "Cause", "Fix"], [
        ["UI shows Not Found", "/v1/demo/audit missing or stale image", "docker-compose up --build -d"],
        ["UI shows 404", "Old process on port 8000", "Run restart.bat, use localhost:8000"],
        ["403 Forbidden", "Secret mismatch", "Align .env with GitHub webhook secret"],
        ["Worker not processing", "Redis down or wrong queue", "docker-compose ps, check agentaudit queue"],
        ["Docker won't start", "Docker Desktop off", "Start Docker Desktop first"],
        ["Stripe example passes", "Demo key too short", "Use 24+ chars after sk_live_"],
    ])

    # Section 11
    add_heading(doc, "11. Roadmap & Extensions", 1)
    add_table(doc, ["Version", "Feature"], [
        ["v2.1", "Admin UI for policy management"],
        ["v2.2", "Custom rule packs per tenant"],
        ["v2.3", "SARIF export for GitHub Advanced Security"],
        ["v2.4", "HashiCorp Vault secret injection"],
        ["v2.5", "Jira / ServiceNow incident automation"],
    ])

    # Section 12
    doc.add_page_break()
    add_heading(doc, "12. Technology Stack & Rationale", 1)
    add_body(doc, "This section explains every major technology used in AgentAuditAI and the engineering reasons behind each choice.")

    add_heading(doc, "Core Language — Python 3.12", 2)
    add_table(doc, ["Aspect", "Detail"], [
        ["Used for", "API gateway, audit engine, workers, configuration"],
        ["Why chosen", "Mature security tooling ecosystem, fast development, strong regex/HMAC support"],
        ["Enterprise fit", "Easy hiring, SIEM/vault/CI integration, consistent Docker deployment"],
    ])

    add_heading(doc, "API Layer", 2)
    add_table(doc, ["Technology", "Used For", "Why Chosen"], [
        ["FastAPI", "Webhook gateway, demo UI, health, docs", "Async performance, validation, built-in Swagger"],
        ["Uvicorn", "ASGI server", "Lightweight, production-ready FastAPI hosting"],
    ])

    add_heading(doc, "Configuration Management", 2)
    add_table(doc, ["Technology", "Used For", "Why Chosen"], [
        ["Pydantic v2 + pydantic-settings", "Environment variable loading", "Type-safe validation, SecretStr for secrets"],
        ["python-dotenv", "Local .env loading", "Keeps secrets out of source code"],
    ])

    add_heading(doc, "Security & Audit Engine", 2)
    add_table(doc, ["Technology", "Used For", "Why Chosen"], [
        ["HMAC + hashlib", "GitHub signature verification", "Constant-time compare, no extra dependencies"],
        ["Compiled regex (re)", "Credential pattern detection", "Fast, stateless, auditable rules"],
        ["gc (garbage collector)", "Memory purge after audit", "Enforces no-retention compliance policy"],
    ])

    add_heading(doc, "Background Processing", 2)
    add_table(doc, ["Technology", "Used For", "Why Chosen"], [
        ["Celery", "process_github_webhook_audit task", "Industry-standard async task queue"],
        ["Redis", "Celery broker/backend", "Fast, scalable message broker"],
        ["PostgreSQL + RLS", "organizations + scan_audit_logs", "Multi-tenant isolation at DB layer"],
        ["SQLAlchemy 2.0", "ORM + session management", "Production-grade database access"],
        ["tree-sitter", "Python AST acceleration", "Grammar-aware syntax validation"],
    ])

    add_heading(doc, "Data, Frontend & Deployment", 2)
    add_table(doc, ["Technology", "Used For", "Why Chosen"], [
        ["PostgreSQL", "Tenant orgs + audit metadata with RLS", "Enterprise RDBMS with row-level security"],
        ["HTML + JavaScript", "Demo UI dashboard", "Zero build step, calls /v1/demo/audit"],
        ["Docker + Docker Compose", "postgres + redis + web + worker", "One-command reproducible deployment"],
        ["python-docx + matplotlib", "Word docs with diagrams", "Stakeholder-ready visual deliverables"],
    ])

    add_heading(doc, "Why This Stack Fits Enterprise & Public Companies", 2)
    add_table(doc, ["Organizational Need", "Technology Answer"], [
        ["Fast webhook response", "FastAPI + async + Celery queue"],
        ["Secret verification", "HMAC-SHA256 constant-time compare"],
        ["No diff retention", "Stateless engine + gc.collect()"],
        ["Config safety", "Pydantic Settings + SecretStr"],
        ["Horizontal scalability", "Redis + multiple Celery workers"],
        ["Compliance audit trail", "Metadata-only structured logging"],
        ["Easy deployment", "Docker Compose"],
        ["SIEM / vault / GHES integration", "Modular Python services"],
    ])

    add_heading(doc, "Technologies Intentionally NOT Used", 2)
    add_table(doc, ["Technology", "Why Not Used"], [
        ["Django", "Heavier than needed for a focused webhook API"],
        ["Kafka", "Overkill for current queue volume; Redis is simpler"],
        ["ML / AI scanning", "Regex is faster, explainable, and auditable"],
        ["Storing diffs in database", "Violates data minimization policies"],
        ["React / Vue SPA", "Unnecessary build complexity for demo UI"],
    ])

    add_body(doc, "Summary: AgentAudit AI uses Python + FastAPI + Celery + Redis + PostgreSQL RLS to deliver a fast, compliant, multi-tenant GitHub webhook security platform.")

    doc.add_page_break()
    add_heading(doc, "Appendix A — Environment Template", 1)
    env = doc.add_paragraph(
        "DATABASE_URL=postgresql://agentaudit:agentaudit@postgres:5432/agentaudit\n"
        "REDIS_URL=redis://redis:6379/0\n"
        "GITHUB_WEBHOOK_SECRET=your-32-char-minimum-secret\n"
        "SECRET_KEY=your-long-random-secret-key\n"
        "SECURITY_ENVIRONMENT=development"
    )
    for run in env.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    add_heading(doc, "Appendix B — Contact & Support", 1)
    add_body(doc, "Repository: github.com/mrunalvuppala/github-webhook-audit")
    add_body(doc, "Issues: GitHub Issues tab on the repository")

    return doc


def main() -> None:
    DIAGRAMS_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    images = {
        "architecture": draw_system_architecture(),
        "lifecycle": draw_request_lifecycle(),
        "rls": draw_database_rls(),
        "enterprise": draw_enterprise_deployment(),
        "public": draw_public_company_deployment(),
        "ui": draw_ui_layout(),
        "testing": draw_testing_flow(),
    }

    doc = build_document(images)
    doc.save(OUTPUT)
    doc.save(DESKTOP_OUTPUT)

    print(f"Created: {OUTPUT}")
    print(f"Copied:  {DESKTOP_OUTPUT}")
    print("Diagrams:")
    for key, path in images.items():
        print(f"  - {key}: {path}")


if __name__ == "__main__":
    main()
